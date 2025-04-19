import streamlit as st
import openai
import pandas as pd
from io import StringIO, BytesIO
from docx import Document
import os
import csv
from datetime import datetime

# --- Setup ---
st.set_page_config(page_title="AI Marking Assistant", layout="wide")
st.title("📘 AI-Assisted Assessment Marking Tool")

# --- API Key ---
openai.api_key = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else os.getenv("OPENAI_API_KEY")

# --- Upload Section ---
st.header("1. Upload Assessment Files")
assessment_descriptor = st.file_uploader("Assessment Descriptor (DOCX or TXT)", type=["docx", "txt"])
marking_rubric = st.file_uploader("Marking Rubric (CSV or Excel)", type=["csv", "xlsx"])
minimum_requirements = st.text_area("Minimum Requirements (optional)")

sample_col1, sample_col2, sample_col3 = st.columns(3)
with sample_col1:
    low_sample = st.file_uploader("Low Sample", type=["docx", "txt"], key="low")
with sample_col2:
    medium_sample = st.file_uploader("Medium Sample", type=["docx", "txt"], key="medium")
with sample_col3:
    high_sample = st.file_uploader("High Sample", type=["docx", "txt"], key="high")

# --- Upload Submissions ---
st.header("2. Upload Student Submission")
student_submission = st.file_uploader("Student Work (DOCX or TXT)", type=["docx", "txt"], key="submission")
student_name = st.text_input("Student Name")

# --- Rubric Parser ---
def load_rubric(file):
    if file is not None:
        if file.name.endswith(".csv"):
            return pd.read_csv(file)
        else:
            return pd.read_excel(file)
    return None

rubric_df = load_rubric(marking_rubric)
if rubric_df is not None:
    st.subheader("Rubric Preview")
    st.dataframe(rubric_df)

# --- DOCX Reader ---
def read_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def read_text(file):
    return file.read().decode("utf-8")

# --- AI Marking (Updated for openai>=1.0.0 and using GPT-3.5) ---
def generate_feedback(text, rubric, sample_level="medium"):
    prompt = f"""
You are an experienced academic marker. Based on the following rubric and sample level ({sample_level}), mark the student's work below.
Rubric:
{rubric.to_string(index=False)}

Student Work:
{text}

Please provide:
1. A mark for each criterion (mention the criterion).
2. Constructive comments per criterion.
3. An overall summary of the student's performance.
"""
    try:
        client = openai.OpenAI()
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful academic marking assistant."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generating feedback: {e}"

# --- Process Submission ---
if st.button("Generate AI Feedback"):
    if student_submission and rubric_df is not None and student_name:
        with st.spinner("Reading and analysing the student work..."):
            if student_submission.name.endswith(".docx"):
                student_text = read_docx(student_submission)
            else:
                student_text = read_text(student_submission)

            result = generate_feedback(student_text, rubric_df)
            st.subheader("AI Feedback")
            feedback_text = st.text_area("Generated Feedback (editable)", result, height=400)

            # Optional manual scoring
            st.subheader("Manual Score Adjustments (Optional)")
            scores = []
            for i, row in rubric_df.iterrows():
                score = st.slider(f"{row[0]}", 0, 10, 5)
                scores.append((row[0], score))

            # Export section
            st.subheader("Export Results")
            if st.button("Download as CSV"):
                output = StringIO()
                writer = csv.writer(output)
                writer.writerow(["Student Name", "Criterion", "Score"])
                for criterion, score in scores:
                    writer.writerow([student_name, criterion, score])
                st.download_button(
                    label="Download Scores CSV",
                    data=output.getvalue(),
                    file_name=f"{student_name}_scores.csv",
                    mime="text/csv"
                )

            if st.button("Download Feedback as DOCX"):
                doc = Document()
                doc.add_heading(f"Assessment Feedback for {student_name}", level=1)
                doc.add_paragraph(feedback_text)
                output_docx = BytesIO()
                doc.save(output_docx)
                st.download_button(
                    label="Download Feedback DOCX",
                    data=output_docx.getvalue(),
                    file_name=f"{student_name}_feedback.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    else:
        st.warning("Please upload a student submission, rubric, and provide the student's name.")
